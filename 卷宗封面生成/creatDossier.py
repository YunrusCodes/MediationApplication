import os
import shutil
import sys
if getattr(sys, 'frozen', False):
    dir_path = os.path.dirname(sys.executable)
else:
    dir_path = os.path.dirname(os.path.abspath(__file__))

# 新增建立 output 資料夾的功能
output_folder = os.path.join(dir_path,'生成卷宗封面')
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

import pymysql
from sqlalchemy import create_engine
import re

# 資料庫連線設定
host = 'localhost'
port = 3306
user = sys.argv[1] if len(sys.argv) >= 2 else input("請輸入使用者名稱：")
password = sys.argv[2] if len(sys.argv) >= 2 else input("請輸入密碼：")
database = 'mediation'
try:
    # 連接到 MySQL 資料庫
    conn = pymysql.connect(host=host, port=port, user=user, password=password,database=database)
    cursor = conn.cursor()
    # 建立 SQLAlchemy 引擎
    engine = create_engine(f'mysql+pymysql://{user}:{password}@{host}:{port}/{database}')
except Exception as e:
   input(f"發生錯誤:{e}\n請按任一鍵離開")
   sys.exit()

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import datetime


def setKaiTi(font,size,bold):
  font.name = '標楷體'
  font.bold = bold
  r = font._element.rPr  # 取得 w:rPr 元素
  rFonts = docx.oxml.OxmlElement('w:rFonts')  # 新增 w:rFonts 元素
  rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')  # 設定 w:eastAsia 屬性為標楷體
  r.append(rFonts)  # 將 w:rFonts 元素加入 w:rPr 元素
  font.size = Pt(size)

def replace_placeholder(cell, placeholder, value, size, centered=False):
    if placeholder in cell.text:
        cell.text = value if value is not None else ''
        original_font = cell.paragraphs[0].runs[0].font
        setKaiTi(original_font, size, True)
        if centered:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

try:
    # 輸入收件編號
    hanndler = input("請輸入經辦人姓名或輸入 exit 離開：\n")
    input_num = input("請輸入收件編號或輸入 exit 離開：\n") if hanndler != 'exit' else 'exit'
    while input_num != 'exit' and hanndler != 'exit':
        data_1 = {}
        data_2 = {}
        table = 'med_case'
        # 查詢收件編號對應的資料
        cursor.execute(f"SELECT * FROM {database}.{table} WHERE `收件編號`='{input_num}'")
        result = cursor.fetchone()
        # 判斷是否有查詢到資料
        if result:
            # 將查詢到的資料放入字典
            data_1 = {
                '收件編號': result[0],
                '收件日期': result[1],
                '案號': result[2],
                '聲請人': result[3],
                '對造人': result[4],
                '事由': result[5],
                '過程摘要': result[6],
            }
            for key, value in data_1.items():
                print(f"{key}: {value}")
            table = 'contents'
            cursor.execute(f"SELECT * FROM {table} WHERE `案號`='{data_1['案號']}'")
            result = cursor.fetchone()
            if result:
                # 將查詢到的資料放入字典
                data_2 = {
                    '收件編號': result[0],
                    '案號': result[1],
                    '調解案件轉介單〈函〉': result[2],
                    '轉介單位': result[3],
                    '轉介人': result[4],
                    '聲請調解書': result[5],
                    '委任書或請求權讓與同意書': result[6],
                    '委任人': result[7],
                    '受任人': result[8],
                    '身分證或戶籍謄本資料〈營利事業登記證影本〉': result[9],
                    '調解事件處理單': result[10],
                    '證據文件資料': result[11],
                    '調解期日通知書副本': result[12],
                    '調解期日通知書送達證書': result[13],
                    '調解筆錄': result[14],
                    '函報法院審核函副本': result[15],
                    '報院審查日期': result[16],
                    '法院核定函〈補正函〉': result[17],
                    '法院發還日期': result[18],
                    '法院審查結果': result[19],
                    '調解書': result[20],
                    '調解委員': result[21],
                    '主席': result[22],
                    '檢送調解書予當事人函副本': result[23],
                    '調解書送達證書': result[24],
                    '最後送達日期':result[25],
                    '發給調解不成立證明聲請書': result[26],
                    '調解不成立證明書': result[27],
                    '刑事事件調解不成立移送偵查聲請書': result[28],    
                    '刑事事件調解不成立移送偵查書副本': result[29],
                    '調解撤回聲請書': result[30],
                    '其它(底頁)': result[31],
                    '私下和解註記': result[32],
                    '報院審查文號': result[33],
                    '法院發還文號': result[34]
                }
                
                for key, value in data_2.items():
                    print(f"{key}: {value}")

                doc_path = os.path.join(dir_path,'空白卷宗封面.docx')
                # 開啟並編輯送達證書
                doc = docx.Document(doc_path)
                unknow = 0
                
                if data_1['案號'] and data_1['案號'] != None:
                    # 新增受通知人名字作為檔名
                    re_num = data_1['收件編號']
                    case_num = data_1['案號'].split('年')[0]+data_1['案號'].split('年')[1].replace('調字第','').replace('號','')
                    print(case_num)
                    doc_name = '卷宗封面(收{:03d}案{}).docx'.format( re_num, case_num)
                    doc_path = os.path.join(output_folder, doc_name)
                    # 複製空白送達證書
                    shutil.copy(os.path.join(dir_path, '空白卷宗封面.docx'), doc_path)
                    # 開啟並編輯送達證書
                    doc = docx.Document(doc_path)
                else:
                    doc_name = '卷宗封面({}).docx'.format(f'未知案號{str(unknow)}')
                    unknow += 1
                    doc_path = os.path.join(output_folder, doc_name)
                    # 複製空白送達證書
                    shutil.copy(os.path.join(dir_path, '空白卷宗封面.docx'), doc_path)
                    # 開啟並編輯送達證書
                    doc = docx.Document(doc_path)

                # 設定受通知人內容與字型
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            placeholders = {
                                '收件編號：': (f'收件編號：{str(data_1["收件編號"]).zfill(3)}號', 14,False),
                                '(案號填入處)': (data_1['案號'], 16,True),
                                '(案由填入處)': (data_1['事由'], 14,False),
                                '(收案日填入處)': (f"{data_1['收件日期'].split('日')[0]}日", 14,False),
                                '(結案日填入處)': (data_2['最後送達日期'], 14,False),
                                '(聲請人填入處)': (data_1['聲請人'].replace(',', '、'), 14,False),
                                '(對造人填入處)': (data_1['對造人'].replace(',', '、'), 14,False),
                                '(調解結果填入處)': ( '私下和解' if data_2['私下和解註記'] != None and data_2['私下和解註記'].strip() != None else '調解不成立' if data_2['調解書'] == 0 else '調解成立', 14, True),
                                '(審核結果填入處)': (data_2['法院審查結果'].replace(',', '、'), 14,True),
                                '(保存年限填入處)': (f"{15 if data_2['調解書'] != 0 else 3} 年", 14,True),
                                '(轉介機構填入處)': (data_2['轉介單位'], 12,True),
                                '(轉介人姓名填入處)': (data_2['轉介人'], 14,True),
                                '(主席姓名填入處)': ('\n'+data_2['主席'], 14,True),
                                '(調解委員姓名填入處)': ('\n'+data_2['調解委員'], 14,True),                            
                                '(經辦人姓名填入處)' : ('\n'+hanndler, 14,True),
                            }

                            today = datetime.date.today()
                            year, month, day = today.year, today.month, today.day
                            placeholders.update({
                                '(歸檔日期填入處)': (f'中華民國 {year-1911} 年 {month} 月 {day} 日', 14,False),
                                '(保存始期填入處)': (f'中華民國 {year-1911} 年 {month} 月 {day} 日', 14,False),
                                '(保存終期填入處)': (f'中華民國 {year-1911 + (15 if data_2["調解書"] != None else 3)} 年 {month} 月 {day} 日', 14,False)
                            })

                            apps = data_2.get('委任人', '').replace(',', '、').split('、') if data_2.get('委任人') else []
                            accs = data_2.get('受任人', '').replace(',', '、').split('、') if data_2.get('受任人') else []

                            applicant_sub = [accs[i] for i in range(len(apps)) if apps[i] in data_1['聲請人']]
                            opponent_sub = [accs[i] for i in range(len(apps)) if apps[i] in data_1['對造人']]
                            placeholders.update({
                                '(聲請人代理人填入處)': ('、'.join(applicant_sub) if applicant_sub else '', 14,False),
                                '(對造人代理人填入處)': ('、'.join(opponent_sub) if opponent_sub else '', 14,False),
                            })

                            for placeholder, (value,size,centered) in placeholders.items():
                                replace_placeholder(cell, placeholder, value, size, centered = centered)

                # 將文件中所有出現"收件編號："的地方的該文字設為紅色並設定為標楷體
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if '收件編號：' in cell.text:
                                for paragraph in cell.paragraphs:
                                    if '收件編號：' in paragraph.text:
                                        # 將段落文字依照 "收件編號：" 分開，分成三個 runs 分別顯示前、中、後的文字
                                        pre_text, keyword, post_text = paragraph.text.partition('收件編號：')
                                        paragraph.text = ''
                                        key_run = paragraph.add_run(keyword)
                                        post_run = paragraph.add_run(post_text)
                                        # 設定中間的 run 的文字為紅色並設定為標楷體
                                        key_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                                        setKaiTi(key_run.font, 14, True)  # 設定為標楷體
                                        setKaiTi(post_run.font, 14, True)

                contents = ['調解案件轉介單〈函〉' ,'聲請調解書' ,'委任書或請求權讓與同意書' ,'身分證或戶籍謄本資料〈營利事業登記證影本〉', '調解事件處理單'
                            ,'證據文件資料' ,'調解期日通知書副本' ,'調解期日通知書送達證書', '調解筆錄'
                            ,'函報法院審核函副本' ,'法院核定函〈補正函〉'  ,'調解書' 
                            ,'檢送調解書予當事人函副本', '調解書送達證書', '發給調解不成立證明聲請書' 
                            ,'調解不成立證明書', '刑事事件調解不成立移送偵查聲請書' 
                            ,'刑事事件調解不成立移送偵查書副本', '調解撤回聲請書', '其它(底頁)' ]
                new_contents = []
                for content in contents:
                    if data_2[content] != 0:
                        new_contents.append(content)
                contents = new_contents
                
                pageCounter = 1
                for table in doc.tables:
                    for row in table.rows:
                        if row.cells[1].text in contents :
                            row.cells[2].text = str(pageCounter)
                            row.cells[4].text = str(pageCounter + data_2[row.cells[1].text] - 1)
                            pageCounter += data_2[row.cells[1].text]
                            row.cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            row.cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            setKaiTi(row.cells[2].paragraphs[0].runs[0].font, 14, True)
                            setKaiTi(row.cells[4].paragraphs[0].runs[0].font, 14, True)

                doc.save(doc_path)
            else:
              print('查無相關目錄資料: ', data_1['案號'],'\n請確認目錄之案號是否正確上傳')
        else:
            print('無效的輸入')

        input_num = input("請輸入收件編號或輸入 exit 離開：\n")
    print('歡迎再次使用')
except Exception as e:
   input(f"發生錯誤:{e}\n請按任一鍵離開")
   sys.exit()