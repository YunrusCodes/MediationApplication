import sys
import os
arg = None
arg_year = None

try:
  arg_year = sys.argv[1]
  arg = sys.argv[2]
  if arg :
    print("取得引數:",arg)
  else:
    print("未傳入引數，採取優先取代方法")  

  print("年份: ",arg_year)
except:
  print("未傳入引數，採取優先取代方法")
  if not arg_year:
    arg_year = input("請輸入民國年份:")
  print("年份: ",arg_year)

if getattr(sys, 'frozen', False):
    absPath = os.path.dirname(sys.executable)
else:
    absPath = os.path.dirname(os.path.abspath(__file__))

import docx
import re
import unicodedata
def to_halfwidth(s):
    # 把全形字轉成半形字
    return unicodedata.normalize('NFKC', s)

def to_chinese_numeral(num):
    chinese_numerals = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九']
    chinese_unit_numerals = ['', '十', '百', '千', '萬', '十', '百', '千', '億', '十', '百', '千', '兆']
    # 先將數字轉成字串
    num_str = str(num)
    # 將每一個數字轉換成中文數字
    num_chinese = ''.join([chinese_numerals[int(digit)] for digit in num_str])
    # 將中文數字加上單位
    result = ''
    for i in range(len(num_chinese)):
        result += chinese_unit_numerals[len(num_str) - i - 1]
        result += num_chinese[i]
    # 處理特殊情況
    if num_str[-2:] == '10':
        result = result[:-2] + chinese_unit_numerals[1]
    return result.replace('零萬', '萬').replace('零億', '億').replace('零兆', '兆').replace('零零', '零').replace('零十', '零')

def count_lines(s, w):
    line_count = 0
    word_count = 0
    for char in s:
        # 判斷是否為中文字或中文標點符號
        if '\u4e00' <= char <= '\u9fff' or char in '，。！？；：、（）『』「」《》【】':
            word_count += 1
        # 判斷是否為數字或其他非中文符號
        elif char.isdigit() or char in '.,;?!':
            word_count += 0.5
        # 判斷是否為換行符號
        elif char == '\n':
            line_count += 1
            word_count = 0
        # 忽略其他非中文字符
        else:
            pass
        # 如果超過指定的字數，換行計數
        if word_count > w:
            line_count += 1
            word_count = 0
    # 如果還有未滿一行的字，計為一行
    if word_count > 0:
        line_count += 1
    return line_count


class Case:
    def __init__(self, file_path):
        self.file_path = file_path
        self.number = None             #收件編號
        self.time = None               #收件日期
        self.case_number = None        #案號
        self.applicants = []           #聲請人
        self.opponents = []            #對造人
        self.case_reason = None        #事由
        self.summary = ""            #過程摘要
        self.mediation_result = None   #調解結果

    def parse_record(self): # 調解筆錄
        # 打開Word文件
        doc = docx.Document(self.file_path)
        # 提取表格中的內容
        tables = doc.tables
        for table in tables:
            for row in table.rows:
                row_data = []
                try:
                    for i, data in enumerate(row_data): # 消除空白
                        row_data[i] = data.replace(" ", "").replace("\n", "").replace("\t", "")

                    for cell in row.cells:
                        row_data.append(cell.text)
                    # print(row_data)
                    # 解析表格資料
                    if re.search(r'^聲請人\d*$', row_data[0]):
                        self.applicants.append(row_data[1].replace(" ","").replace("\n","、"))
                    if re.search(r'^對造人\d*$', row_data[0]):
                        self.opponents.append(row_data[1].replace(" ","").replace("\n","、"))

                    for datas in row_data:
                        match = re.search(r'收件編號：(\d+)', datas)  # 收件編號
                        if match and self.number is None:
                            self.number = match.group(1).strip() 
                        match = re.search(r'案號(\d+年\S*調字第\d+號)\s*', datas.replace(" ","")) #案號
                        if match and self.case_number is None:
                            self.case_number = match.group(1).strip()
                except IndexError:
                  print(f"Error: row_data index out of range. row_data: {row_data}")


    def parse_petition(self):  # 聲請調解書
        # 打開Word文件
        doc = docx.Document(self.file_path)
        # 提取表格中的內容
        tables = doc.tables
        for table in tables:
            for row in table.rows:
                row_data = []

                try:
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                        # print(row_data)
                        # 解析表格資料
                        # for i, data in enumerate(row_data):
                        #     row_data[i] = data.replace(" ", "").replace("\n", "").replace("\t", "")

                    if len(row_data) > 2 and row_data[0] == '聲請人' and not row_data[1] in self.applicants:
                        self.applicants.append(row_data[1])
                    if len(row_data) > 2 and row_data[0] == '對造人' and not row_data[1] in self.opponents :
                        self.opponents.append(row_data[1])
    
                    for datas in row_data:
                        match = re.search(r'收件編號：(\d+)', datas)  # 收件編號
                        if match and self.number is None:
                            self.number = match.group(1).strip()

                        match = re.search(r'收件日期：(\d+年\d+月\d+日\d+時\d+分)', datas) #收件日期
                        if match :
                            self.time = match.group(1).strip()
                        
                        if '案號' in datas and '調字第' in datas :
                            self.case_number = datas.replace('案號','')
                        match = re.search(r'上開當事人間因「(.*)」.*聲請調解.*', datas) #事由
                        if match :
                            self.case_reason = match.group(1)
                            self.case_reason = self.case_reason.replace("事件","")
                except IndexError:
                      print(f"Error: row_data index out of range. row_data: {row_data} from",self.file_path)


    def parse_meditation(self): #調解書
        # 打開Word文件
        doc = docx.Document(self.file_path)
        # 提取表格中的內容
        tables = doc.tables
        for table in tables:
            for row in table.rows:
                row_data = []
                try:
                    for cell in row.cells:
                        row_data.append(cell.text)
                    for i, data in enumerate(row_data): # 消除空白
                        row_data[i] = data.replace(" ", "").replace("\n", "").replace("\t", "")
                    for datas in row_data:
                        match = re.search(r'收件編號：(\d+)', datas)  # 收件編號
                        if match and self.number is None:
                            self.number = match.group(1).strip() 
                        match = re.search(r'案號(\d+年\S*調字第\d+號)\s*', datas.replace(" ","")) #案號
                        if match and self.case_number is None:
                            self.case_number = match.group(1).strip()
                        

                except IndexError:
                  print(f"Error: row_data index out of range. row_data: {row_data}")

        index = 0
        
        # #上開當事人間因「解除套繪管制糾紛事件」，於民國112年01月10日9時10分在本會調解室調解成立，內容如下：
        # date = doc.paragraphs[0].text
        # year = date.split('民國')[1].split('年')[0]
        # month = date.split('年')[1].split('月')[0]
        # day = date.split('月')[1].split('日')[0]
        # date = year + '.' + month + '.' + day

        # Summary = ""
        # for i, para in enumerate(doc.paragraphs):
        #     if "雙方同意如下" in para.text:
        #         index = i  # 更新 index 變數的值
        #         break
        # for i in range(index + 1, len(doc.paragraphs)):
        #     Summary += doc.paragraphs[i].text
            
        # sumaNum = 1
        # Summary = Summary.replace('\n','').replace('。','。\n')
        
        # for summaryList in Summary.split('。\n'):
        #     summaryList = summaryList.replace(" ","")
        #     if '本件現正在臺灣地方檢察署偵查審理中' in summaryList:
        #       break
        #     elif summaryList:  
        #       if to_chinese_numeral(sumaNum )+ '、' in summaryList:
        #         self.summary += summaryList + '\n'
        #       else:
        #         self.summary += to_chinese_numeral(sumaNum )+ '、' + summaryList + '\n'
        #       sumaNum += 1
        # print(date)
        # self.summary = date + '\n' + Summary
        # if self.summary:
        #    self.summary = to_halfwidth(self.summary)

#------------------------------------------------------------------------------------
        #上開當事人間因「解除套繪管制糾紛事件」，於民國112年01月10日9時10分在本會調解室調解成立，內容如下：
        date = doc.paragraphs[0].text
        year = date.split('民國')[1].split('年')[0]
        month = date.split('年')[1].split('月')[0]
        day = date.split('月')[1].split('日')[0]
        date = year + '.' + month + '.' + day
        index = 0
        for i, para in enumerate(doc.paragraphs):
            if "雙方同意如下" in para.text:
                index = i  # 更新 index 變數的值
                break
        Summary = ""
        for i in range(index + 1, len(doc.paragraphs)):
            Summary += doc.paragraphs[i].text
            
        sumaNum = 1
        Summary = Summary.replace('\n','').replace('。','。\n')
        
        for summaryList in Summary.split('。\n'):
            summaryList = summaryList.replace(" ","")
            if '本件現正在臺灣地方檢察署偵查審理中' in summaryList:
              break
            elif summaryList:  
              if to_chinese_numeral(sumaNum )+ '、' in summaryList:
                self.summary += summaryList + '\n'
              else:
                self.summary += to_chinese_numeral(sumaNum )+ '、' + summaryList + '\n'
              sumaNum += 1
        self.summary = date + '\n' + self.summary
        if self.summary:
           self.summary = to_halfwidth(self.summary)
#------------------------------------------------------------------------------------

def update_case_data(existing_case, new_case):
    if existing_case == new_case:
        return existing_case
    updated_case = Case(existing_case.file_path)
    updated_case.number = existing_case.number if existing_case.number is not None else new_case.number
    updated_case.time = existing_case.time if existing_case.time is not None else new_case.time
    updated_case.case_number = existing_case.case_number if existing_case.case_number is not None else new_case.case_number
    updated_case.applicants = existing_case.applicants if existing_case.applicants else new_case.applicants
    updated_case.opponents = existing_case.opponents if existing_case.opponents else new_case.opponents
    updated_case.case_reason = existing_case.case_reason if existing_case.case_reason is not None else new_case.case_reason
    updated_case.summary = existing_case.summary if existing_case.summary.strip() else new_case.summary
    if existing_case.applicants and new_case.applicants:
        updated_case.applicants = list(set(existing_case.applicants + new_case.applicants))
    if existing_case.opponents and new_case.opponents:
        updated_case.opponents = list(set(existing_case.opponents + new_case.opponents))
    return updated_case

cases = {}
def search_and_parse_files(file_type, pattern, parse_func):
    global cases
    path1 = os.path.join(os.path.dirname(absPath), '格式轉換', '請將欲處理文件備份後放入此處')
    files1 = [os.path.join(root, file) for root, dirs, filenames in os.walk(path1)
              for file in filenames if re.search(pattern, file) and os.path.splitext(file)[1] == ".docx"]

    path2 = os.path.join(os.path.dirname(absPath), '通知書生成', '通知書輸出於此')
    files2 = [os.path.join(root, file) for root, dirs, filenames in os.walk(path2)
              for file in filenames if re.search(pattern, file) and os.path.splitext(file)[1] == ".docx"]
    print(files1 + files2)
    for file in files1 + files2:
        case = Case(file)
        parse_func(case)
        if case.case_number is not None and case.case_number.startswith(arg_year):
            if str(case.case_number) in cases:
                existing_case = cases[str(case.case_number)]
                updated_case = update_case_data(existing_case, case)
                cases[str(case.case_number)] = updated_case
            else:
                cases[str(case.case_number)] = case


try:
    if arg == "解析聲請調解書":
    # 搜尋並解析聲請調解書
       search_and_parse_files("petition", "聲請調解書", Case.parse_petition)
    elif arg == "解析調解筆錄":
    # 搜尋並解析調解筆錄
       search_and_parse_files("record", r'^調解筆錄(?:\([\w\s]+\))?\.docx$|^\\d+\\.[\\s\\S]*調解筆錄[\\s\\S]*\.docx$', Case.parse_record)
    elif arg == "解析調解書":
    # 搜尋並解析調解書
       search_and_parse_files(r"mediation", r"^(?!.*送達)(?!.*聲請)(?:\d+\.[\s\S]*調解書[\s\S]*|調解書(?:\([\w\s]+\))?)\.docx$", Case.parse_meditation)
    else:
       search_and_parse_files("petition", "聲請調解書", Case.parse_petition)  
       search_and_parse_files("record", r'^調解筆錄(?:\([\w\s]+\))?\.docx$|^\\d+\\.[\\s\\S]*調解筆錄[\\s\\S]*\.docx$', Case.parse_record)
       search_and_parse_files(r"mediation", r"^(?!.*送達)(?!.*聲請)(?:\d+\.[\s\S]*調解書[\s\S]*|調解書(?:\([\w\s]+\))?)\.docx$", Case.parse_meditation)

    import os
    import pandas as pd

    # 輸出檔案的路徑
    path = os.path.join(absPath, '2.解析結果輸出於此', '資訊列表.xlsx')

    if os.path.isfile(path):
        # 讀取現有的資料
        existing_df = pd.read_excel(path, dtype={"收件編號": str})
        
        # 新增新資料
        new_data = []
        for key, value in cases.items():
            new_data.append({
                "收件編號": value.number,
                "收件日期": value.time,
                "案號": value.case_number,
                "聲請人": ", ".join(value.applicants),
                "對造人": ", ".join(value.opponents),
                "事由": value.case_reason,
                "過程摘要": value.summary
            })
        new_df = pd.DataFrame(new_data, columns = ['收件編號', '收件日期', '案號', '聲請人', '對造人', '事由', '過程摘要'])

        # 合併現有資料與新資料
        merged_df = pd.concat([existing_df, new_df], ignore_index=True)
        
        # 將 DataFrame 存成 xlsx 檔案
        with pd.ExcelWriter(path) as writer:
            merged_df.to_excel(writer, index=False)
    else:
        # 如果檔案不存在，直接存新資料
        data = []
        for key, value in cases.items():
            data.append({
                "收件編號": value.number,
                "收件日期": value.time,
                "案號": value.case_number,
                "聲請人": ", ".join(value.applicants),
                "對造人": ", ".join(value.opponents),
                "事由": value.case_reason,
                "過程摘要": value.summary
            })
        df = pd.DataFrame(data, columns = ['收件編號', '收件日期', '案號', '聲請人', '對造人', '事由', '過程摘要'])
        # 將 DataFrame 存成 xlsx 檔案
        with pd.ExcelWriter(path) as writer:
            df.to_excel(writer, index=False)
except Exception as e:
   input(f"發生錯誤:{e}\n請按任一鍵離開")
   sys.exit()
