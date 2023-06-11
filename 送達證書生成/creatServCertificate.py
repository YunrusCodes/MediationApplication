import os
import openpyxl
import shutil
def getTwRegion(address):
   if address == None:
      return None
   if '鄉' in address: 
     return address.split('鄉')[0] + '鄉'
   if '鎮' in address: 
     return address.split('鎮')[0] + '鎮'
   if '區' in address: 
     return address.split('區')[0] + '區'
   if '島' in address: 
     return address.split('島')[0] + '島'
   if '釣魚臺釣魚臺' in address: 
     return '釣魚臺釣魚臺'
   if '釣魚台釣魚台' in address: 
     return '釣魚台釣魚台'
   if '釣魚臺' in address: 
     return address.split('釣魚臺')[0] + '釣魚臺'
   if '釣魚台' in address: 
     return address.split('釣魚台')[0] + '釣魚台'
   if '市' in address: 
     return address.split('市')[0] + '市'

import sys
if getattr(sys, 'frozen', False):
    dir_path = os.path.dirname(sys.executable)
else:
    dir_path = os.path.dirname(os.path.abspath(__file__))
    
# 讀取檔案
info_book = openpyxl.load_workbook(os.path.join(dir_path, "列出已送達案件編號.xlsx"))

# 新增建立 output 資料夾的功能
output_folder = os.path.join(dir_path,'待印送達證書')
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# 從info_book中讀取資訊並填入workbook中
info_sheet = info_book[info_book.sheetnames[0]]

# 創建郵遞區號表的字典
postalCode_dict = {}
for row in range(2, info_sheet.max_row):
    region = info_sheet.cell(row=row, column=6).value # 行政地區
    if not region:
      break
    postalCode = info_sheet.cell(row=row, column=7).value # 郵遞區號
    postalCode_dict[region] = postalCode
    if '臺' in region: # 台南市、台北市
      region = region.replace('臺','台')
      postalCode_dict[region] = postalCode

doc_type = info_sheet.cell(row=1, column=4).value # J1儲存格

cases_num = []
for row in range(2, info_sheet.max_row):
   case_num = info_sheet.cell(row=row, column=1).value
   if case_num != None:
      cases_num.append(case_num)

info_book.close()

import pymysql
from sqlalchemy import create_engine
import re
import sys
# 資料庫連線設定
host = 'localhost'
port = 3306
user = sys.argv[1] if len(sys.argv) >= 2 else input("請輸入使用者名稱：")
password = sys.argv[2] if len(sys.argv) >= 2 else input("請輸入密碼：")
database = 'mediation'
table = 'participant'
# 連接到 MySQL 資料庫
try:
  conn = pymysql.connect(host=host, port=port, user=user, password=password,database = database)
  cursor = conn.cursor()
  # 建立 SQLAlchemy 引擎
  engine = create_engine(f'mysql+pymysql://{user}:{password}@{host}:{port}/{database}')

  receivers = []
  # 從 MySQL 資料庫中取得涉及到某案件編號的所有參與人的資料
  query = f"SELECT participant_name, case_num, address FROM {table} WHERE case_num = %s"
  for case_num in cases_num:
    if cursor.execute(query, (case_num,)) > 0:
      receivers.extend(cursor.fetchall())
    elif cursor.execute(query, (re.sub(r'(\d+)', lambda x: x.group(1).zfill(3), case_num),)) > 0:
      receivers.extend(cursor.fetchall())
    else:
      print(f"查無資料{case_num}")
    conn.commit()

  import docx
  from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
  from docx.oxml import OxmlElement
  from docx.oxml.ns import qn
  from docx.shared import Pt

  for recipient, caseNum, address in receivers:
      if None in (recipient, caseNum, address):
          print("此組資料異常:")
          print(f"收件人:{recipient} 案號:{caseNum} 地址:{address}")
          if recipient is None:
              recipient = input("請補齊收件人，或直接案enter跳過處裡:\n")
              if not recipient:
                print('資料異常，跳過處裡')
                continue
              else:
                cursor.execute("use mediation")
                cursor.execute(f"UPDATE participant SET participant_name = '{recipient}' WHERE case_num = '{caseNum}' OR address = '{address}'")
                conn.commit()
          if caseNum is None:
              caseNum = input("請補齊案號，或直接案enter跳過處裡:\n")
              if not caseNum:
                print('資料異常，跳過處裡')
                continue
              else:
                cursor.execute("use mediation")
                cursor.execute(f"UPDATE participant SET case_num = '{caseNum}' WHERE participant_name = '{recipient}' OR address = '{address}'")
                conn.commit()
          if address is None:
              address = input("請補齊地址，或直接案enter跳過處裡:\n")
              if not address:
                print('資料異常，跳過處裡')
                continue
              else:
                cursor.execute("use mediation")
                cursor.execute(f"UPDATE participant SET address = '{address}' WHERE participant_name = '{recipient}' AND case_num = '{caseNum}'")
                conn.commit()

      try:
        address = str(postalCode_dict[getTwRegion(address)]) + address #郵遞區號+地址
      except KeyError:
        address = input(f"錯誤地址請更正:{address}\n")
        cursor.execute("use mediation")
        cursor.execute(f"UPDATE participant SET address = '{address}' WHERE participant_name = '{recipient}' AND case_num = '{caseNum}'")
        conn.commit()
        try:
          address =  str(postalCode_dict[getTwRegion(address)]) + address
        except Exception:
           print('異常的地址，請於生成文件中修正資料:')
           address =  address
      
      if recipient and recipient != None:
          # 新增受通知人名字作為檔名
          doc_name = '送達證書({}).docx'.format(recipient)
          doc_path = os.path.join(output_folder, doc_name)
          # 複製空白送達證書
          shutil.copy(os.path.join(dir_path, '空白送達證書.docx'), doc_path)
          # 開啟並編輯送達證書
          doc = docx.Document(doc_path)


          # 設定受通知人內容與字型
          doc.tables[0].rows[2].cells[1].text = recipient
          font = doc.tables[0].rows[2].cells[1].paragraphs[0].runs[0].font
          font.name = '標楷體'
          r = font._element.rPr  # 取得 w:rPr 元素
          rFonts = docx.oxml.OxmlElement('w:rFonts')  # 新增 w:rFonts 元素
          rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')  # 設定 w:eastAsia 屬性為標楷體
          r.append(rFonts)  # 將 w:rFonts 元素加入 w:rPr 元素
          font.size = Pt(16)
          doc.tables[0].rows[2].cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

          # 設定地址內容與字型
          doc.tables[0].rows[4].cells[1].text = address
          font = doc.tables[0].rows[4].cells[1].paragraphs[0].runs[0].font
          font.name = '標楷體'
          r = font._element.rPr  # 取得 w:rPr 元素
          rFonts = docx.oxml.OxmlElement('w:rFonts')  # 新增 w:rFonts 元素
          rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')  # 設定 w:eastAsia 屬性為標楷體
          r.append(rFonts)  # 將 w:rFonts 元素加入 w:rPr 元素
          font.size = Pt(11)

          # 設定案號內容與字型
          caseNum = caseNum.replace(" ","")
          if '民' in caseNum:
            doc.tables[0].rows[6].cells[1].text = caseNum.split('民')[0] + '民' + caseNum.split('民')[1]
          elif '刑' in caseNum:
            doc.tables[0].rows[6].cells[1].text = caseNum.split('刑')[0] + '刑' + caseNum.split('刑')[1]
          font = doc.tables[0].rows[6].cells[1].paragraphs[0].runs[0].font
          font.name = '標楷體'
          font.bold = True
          r = font._element.rPr  # 取得 w:rPr 元素
          rFonts = docx.oxml.OxmlElement('w:rFonts')  # 新增 w:rFonts 元素
          rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')  # 設定 w:eastAsia 屬性為標楷體
          r.append(rFonts)  # 將 w:rFonts 元素加入 w:rPr 元素
          font.size = Pt(14)

          # 設定文件類型內容與字型
          doc.tables[0].rows[8].cells[1].text = doc_type
          font = doc.tables[0].rows[8].cells[1].paragraphs[0].runs[0].font
          font.name = '標楷體'
          r = font._element.rPr  # 取得 w:rPr 元素
          rFonts = docx.oxml.OxmlElement('w:rFonts')  # 新增 w:rFonts 元素
          rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '標楷體')  # 設定 w:eastAsia 屬性為標楷體
          r.append(rFonts)  # 將 w:rFonts 元素加入 w:rPr 元素
          font.size = Pt(14)
          # doc.tables[0].rows[8].cells[1].paragraphs[0].style.font.bold = True
      doc.save(doc_path)
      # # 將內容置中
      # for para in doc.paragraphs:
      #     para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
except Exception as e:
   input(f"發生錯誤:{e}\n請按任一鍵離開")
   sys.exit()


