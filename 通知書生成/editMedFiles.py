import os
import openpyxl
import shutil
from docx import Document
from docx.oxml import OxmlElement
from lxml.etree import XMLSyntaxError
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import traceback
import sys
import json
if getattr(sys, 'frozen', False):
  dir_path = os.path.dirname(sys.executable)
else:
  dir_path = os.path.dirname(os.path.abspath(__file__))

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
class Case:
    def __init__(self, receive_num, case_num, applicant, opponents, date, path):
        self.ReceiveNum = receive_num
        self.CaseNum = case_num
        self.Applicant = applicant
        self.Opponents = opponents
        self.Date = {'Year' : date[0],
                     'Month' : date[1],
                     'Day' : date[2]
                    }
        self.Path = path

    def to_dict(self):
        return {
            'ReceiveNum':self.ReceiveNum,
            'CaseNum': self.CaseNum,
            'Applicant': self.Applicant,
            'Opponents': self.Opponents,
            'Date': self.Date,  # 遍历列表中的每个 Date 对象
            'Path': self.Path
        } 

def run_set_spacing(run, value: int):
    """Set the font spacing for `run` to `value` in twips.

    A twip is a "twentieth of an imperial point", so 1/1440 in.
    """

    def get_or_add_spacing(rPr):
        # --- check if `w:spacing` child already exists ---
        spacings = rPr.xpath("./w:spacing")
        # --- return that if so ---
        if spacings:
            return spacings[0]
        # --- otherwise create one ---
        spacing = OxmlElement("w:spacing")
        rPr.insert_element_before(
            spacing,
            *(
                "w:w",
                "w:kern",
                "w:position",
                "w:sz",
                "w:szCs",
                "w:highlight",
                "w:u",
                "w:effect",
                "w:bdr",
                "w:shd",
                "w:fitText",
                "w:vertAlign",
                "w:rtl",
                "w:cs",
                "w:em",
                "w:lang",
                "w:eastAsianLayout",
                "w:specVanish",
                "w:oMath",
            ),
        )
        return spacing

    rPr = run._r.get_or_add_rPr()
    spacing = get_or_add_spacing(rPr)
    spacing.set(qn('w:val'), str(value))

def create_custom_style(document, style_name):
    styles = document.styles
    custom_style = styles.add_style(style_name, 1)  # 添加新的样式，级别为1

    # 设置段落样式的格式
    custom_style.paragraph_format.left_indent = Pt(30)  # 左缩进
    custom_style.paragraph_format.first_line_indent = Pt(-30)  # 首行缩进
    custom_style.paragraph_format.tab_stops.add_tab_stop(Pt(18))  # 制表位
    custom_style.paragraph_format.space_before = Pt(12)
    return custom_style

def setFixedRowHeight(cell, row_height):
    # 设置行高
    cell.height = Pt(row_height)

def setFixedLineSpacing(paragraph, line_spacing):
    # 设置段落的行距
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = line_spacing

def setKaiTi(font, size = None, bold = None):
    font.name = '標楷體'
    if bold != None:
      font.bold = bold
    r = font._element.rPr  # 取得 w:rPr 元素
    rFonts = OxmlElement('w:rFonts')  # 新增 w:rFonts 元素
    rFonts.set(qn('w:eastAsia'), '標楷體')  # 設定 w:eastAsia 屬性為標楷體
    r.append(rFonts)  # 將 w:rFonts 元素加入 w:rPr 元素
    if size is not None:
        font.size = Pt(size)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def addNewRow(table, targetRowNum):
    new_row = table.add_row()
    new_row_xml = new_row._tr
    table._tbl.insert(targetRowNum, new_row_xml)
    return new_row

def merge_cells_in_row(table, row_index):
    row = table.rows[row_index]
    merge_start_index = None

    for i, cell in enumerate(row.cells):
        if merge_start_index is None:
            merge_start_index = i
        elif cell.text != row.cells[merge_start_index].text:
            if merge_start_index != i - 1:
                table.cell(row_index, merge_start_index).merge(table.cell(row_index, i - 1))
            merge_start_index = i
        else:
            cell.text = ''

    # 合并最后一组相邻单元格（如果需要）
    if merge_start_index is not None and merge_start_index != len(row.cells) - 1:
        table.cell(row_index, merge_start_index).merge(table.cell(row_index, len(row.cells) - 1))
    return row

try:
    # 讀取檔案
    workbook = openpyxl.load_workbook(os.path.join(dir_path, "通知書(空白).xlsx"))
    info_book = os.path.join(dir_path, "受通知人資料整理於此.xlsx")

    # 新增建立 output 資料夾的功能
    output_folder = os.path.join(dir_path, "通知書輸出於此")
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 創建事件編號表的字典
    event_dict = {}
    df = pd.read_excel(info_book, sheet_name='通知書資料').fillna('')
    event_list = df['事件編號表'].tolist()
    for event in event_list:
        if '-' in event:
         event_dict[event.split("-")[0]] = event.split("-")[1]    
    event_dict[''] = ''

    # 創建兩造同意方案的字典
    df = pd.read_excel(info_book, sheet_name='兩造同意方案').fillna('')
    Description_agreement_dict = df.set_index('方案')['內容'].to_dict()
    Description_agreement_dict[''] = ''


    # 創建事件描述方案的字典
    df = pd.read_excel(info_book, sheet_name='事件描述方案').fillna('')
    Description_event_dict = df.set_index('方案')['內容'].to_dict()
    Description_event_dict[''] = ''

    
    #將資料轉譯成字典main_dict
    df = pd.read_excel(info_book, sheet_name='業務總覽').fillna('')
    postDate = df['發文日期'].values[0]
    attendanceTime = df['應到時間'].values[0]
    organiser = df['經辦人'].values[0]
    chairman = df['主席'].values[0]
    serialFolders = []
    df = pd.read_excel(info_book, sheet_name='通知書資料').fillna('') # 讀取 Excel 檔案
    # 將 DataFrame 轉換為以受通知人為鍵的字典
    notc_dict = df.groupby('受通知人').apply(lambda x: x.drop('受通知人', axis=1).to_dict(orient='records')).to_dict() #同人多案
    df = pd.read_excel(info_book, sheet_name='詳細資料').fillna('')
    # 将DataFrame转换为以受通知人为键的字典
    detail_dict = df.groupby('案號').apply(lambda x: x.drop('案號', axis=1).to_dict(orient='records')).to_dict() #同案多人

    main_dict = {}
    # 重新整合字典，避免重複鍵
    for person, notc_infos in notc_dict.items(): #遍歷所有notc_dict
        for notc_info in notc_infos:          #某個人名涉略哪些案件
            for detail in detail_dict[notc_info['案號']]:
              if detail['受通知人'] == person:
                main_dict[notc_info['案號']+'/'+person] = [ notc_info, detail]  # 添加额外的值到现有键的列表中

    for key, value in main_dict.items():
        print(key,'-->',value)

    fix_dict = []
    for key, value in main_dict.items():
        # 核對收件人與案件關係人
        recipient = key.split('/')[1]
        applicant = value[0]['聲請人']
        opponents = value[0]['對造人']
        while not recipient in applicant and not recipient in opponents:
            print(f'收件人-{recipient} 不在此案中 : 聲請人-{applicant} 對造人-{opponents}')

            userQuest = input('是否跳過處裡?[是:輸入"Y" / 否:鍵入ENTER]\n')
            if userQuest == 'y' or userQuest == 'Y':
                break

            userQuest = input('修改收件人?[是:輸入新名稱 / 否:鍵入ENTER]\n')
            if userQuest != '':
                recipient = userQuest

            userQuest = input('修改聲請人?[是:輸入新名稱 / 否:鍵入ENTER]\n')
            if userQuest != '':
                applicant = userQuest
            
            userQuest = input('修改對造人?[是:輸入新名稱 / 否:鍵入ENTER]\n')
            if userQuest != '':
                opponents = userQuest
            userQuest = ''
            new_key = recipient
            new_value1 = {
                '收件編號': value[0]['收件編號'],
                '收件時間': value[0]['收件時間'],
                '受通知人': recipient,
                '地址': value[0]['地址'],
                '案號': value[0]['案號'],
                '聲請人': applicant,
                '對造人': opponents
            }
            new_value2 = {
                '收件編號': value[1]['收件編號'],
                '受通知人': recipient,
                '性別': value[1]['性別'],
                '生日': value[1]['生日'],
                '身份證字號': value[1]['身份證字號'],
                '電話號碼': value[1]['電話號碼'],
                '肇事時間': value[1]['肇事時間'],
                '肇事地點': value[1]['肇事地點'],
                '車輛': value[1]['車輛'],
                '車號': value[1]['車號'],
                '轉介描述': value[1]['轉介描述'],
                '事件描述方案': value[1]['事件描述方案'],
                '兩造同意方案': value[1]['兩造同意方案'],
                '事件編號': value[0]['事件編號']
            }
            if new_key in applicant or new_key in opponents:
                fix_dict.append((new_key,new_value1,new_value2))

    for new_key,new_value1,new_value2 in fix_dict:
        main_dict[new_key] = [new_value1,new_value2]

    #建立舊案字典
    try:
        with open(os.path.join(dir_path,'oldCaseDict.json'), "r",  encoding="utf-8") as file:
            dict_old_cases = json.load(file)
    except FileNotFoundError:
        # 如果文件不存在，創建一個讀取舊檔案的字典
        dict_old_cases = {}
        old_file_folder = os.path.join( os.path.dirname(dir_path), '格式轉換', '請將欲處理文件備份後放入此處')
        for root, dirs, files in os.walk(old_file_folder):
            for dir in dirs:
                # 1.調解筆錄190(聲請人張振強)(對造人黃品翰)(收：111刑152號)(開調解時間 112 年01月10日（二）上午9時 10 分)-(車禍傷害糾紛案)
                if '調解筆錄' in dir and '聲請人' in dir and '開調解時間' in dir:
                    dir = dir.strip()
                    date = [dir.split('開調解時間')[1].split('年')[0], dir.split('開調解時間')[1].split('月')[0], dir.split('開調解時間')[1].split('日')[0]]
                    case = Case(receive_num=None, case_num=None, applicant=None, opponents=None, date=date, path=None)
                    case.receive_num = dir.split('調解筆錄')[1].split('(')[0]
                    case.CaseNum = dir.split('收：')[1].split(')')[0]
                    case.Applicant = dir.split('聲請人')[1].split(')')[0]
                    dic_opp = {}
                    opponents = dir.split('對造人')[1].split(')')[0].split('、')
                    for i, opp in enumerate(opponents):
                        dic_opp[f'對造人{str(i+1)}'] = opp
                    case.Opponents = dic_opp
                    case.Date = {'Year': date[0], 'Month': date[1], 'Day': date[2]}
                    case.Path = dir
                    dict_old_cases[case.receive_num] = case.to_dict()
           

    # 112年06月06日(星期二)上午09時10分
    med_Date_year = attendanceTime.split('年')[0]
    #attendanceTime.split('年')[1] == 06月06日(星期二)上午09時10分
    med_Date_month = attendanceTime.split('年')[1].split('月')[0]
    #attendanceTime.split('年')[1].attendanceTime.split('月')[1] == 06日(星期二)上午09時10分
    med_Date_day = attendanceTime.split('年')[1].split('月')[1].split('日')[0]

    targetFolder = os.path.join(output_folder,f'{med_Date_year}年度調解業務-{med_Date_year}.{med_Date_month}.{med_Date_day}-{organiser}')
    if not os.path.exists(targetFolder):
        print(f'此資料夾不存在:\n{targetFolder}\n請先「生成通知書」')
        input('按任意鍵離開')
        sys.exit()
        
    shutil.copy(info_book, targetFolder)
    opponents = {}

    for root, dirs, files in os.walk(targetFolder):
        if '調解筆錄' in root and '開調解時間' in root: #處理資料夾中的檔案
            applicant = ''
            for file in files:
                if not '~$' in file:
                    print(file)
                    if '01-1.調解通知書-' in file: 
                        applicant = file.split('.調解通知書-')[1].replace('.xlsx','').strip()
                    elif '01-' in file and '.調解通知書-' in file:
                        opponentNum = int(file.split('.調解通知書-')[0].replace('01-','').strip())-1
                        opponents[f'對造人{opponentNum}'] = file.split('.調解通知書-')[1].replace('.xlsx','').strip()
                    print(opponents)
            print(root)        
            caseNum = root.split('收：')[1].split(')')[0] # 案件編號
            receiveNum = main_dict[caseNum+'/'+applicant][0]['收件編號']  # 收件編號
            address =  main_dict[caseNum+'/'+applicant][0]['地址']  # 地址
            receiveDate =  main_dict[caseNum+'/'+applicant][0]['收件時間']  # 收件編號
            
            complete_case_num = caseNum[:3] + '年' + caseNum[3] + '調字第' + caseNum[-3:] + '號'
            reason = event_dict[main_dict[caseNum+'/'+applicant][0]['事件編號']] 
            # 事件描述內容
            Description_event =  Description_event_dict[main_dict[caseNum+'/'+applicant][1]['事件描述方案']]   
            print(opponents)
            for key, value in opponents.items():
                Description_event = Description_event.replace(f'{{{key}}}', opponents[key])
                Description_event = Description_event.replace(f'{{{key}車號}}', str(main_dict[caseNum+'/'+opponents[key]][1]['車號']))
                Description_event = Description_event.replace(f'{{{key}車輛}}', str(main_dict[caseNum+'/'+opponents[key]][1]['車號']))
            
            
            Description_event = Description_event.format(
                聲請人=applicant,
                肇事時間=main_dict[caseNum+'/'+applicant][1]['肇事時間'],
                聲請人車號=main_dict[caseNum+'/'+applicant][1]['車號'],
                聲請人車輛=main_dict[caseNum+'/'+applicant][1]['車輛'],
                肇事地點=main_dict[caseNum+'/'+applicant][1]['肇事地點'],
                轉介描述=main_dict[caseNum+'/'+applicant][1]['轉介描述']
            )

            # 兩造同意內容
            Description_agreement =  Description_agreement_dict[main_dict[caseNum+'/'+applicant][1]['兩造同意方案']]   # 事件描述
            for key, value in opponents.items():
                Description_agreement = Description_agreement.replace(f'{{{key}}}', opponents[key])
                Description_agreement = Description_agreement.replace(f'{{{key}車號}}', str(main_dict[caseNum+'/'+opponents[key]][1]['車號']))
                Description_agreement = Description_agreement.replace(f'{{{key}車輛}}', str(main_dict[caseNum+'/'+opponents[key]][1]['車號']))
            Description_agreement = Description_agreement.format(
                聲請人=applicant,
                肇事時間=main_dict[caseNum+'/'+applicant][1]['肇事時間'],
                聲請人車號=main_dict[caseNum+'/'+applicant][1]['車號'],
                聲請人車輛=main_dict[caseNum+'/'+applicant][1]['車輛'],
                肇事地點=main_dict[caseNum+'/'+applicant][1]['肇事地點'],
                轉介描述=main_dict[caseNum+'/'+applicant][1]['轉介描述']
            )

            record_dir = root #調解筆錄資料夾
            business_dir = os.path.dirname(record_dir) # 調解業務資料夾
            out_dir = os.path.dirname(business_dir) # 通知書輸出於此
            app_dir = os.path.dirname(out_dir) # 通知書生成

            path_dispal = '02.調解事件處理單.xlsx'
            path_request = '03.聲請調解書(筆錄).docx'
            path_record = '04.調解筆錄.docx'
            path_mediation = '05.調解書.docx'
            readfile_list = [path_request , path_record, path_mediation]            
            this_date = f'{med_Date_year}.{med_Date_month}.{med_Date_day}'
            isNewCase = True
            if caseNum in dict_old_cases and  this_date != dict_old_cases[caseNum]['Date']['Year']+'.'+dict_old_cases[caseNum]['Date']['Month']+'.'+dict_old_cases[caseNum]['Date']['Day']:
                print(f'此案為舊案: {caseNum} 聲請人{applicant}')
                try:
                    old_case_date = dict_old_cases[caseNum]['Date']['Year']+'年'+dict_old_cases[caseNum]['Date']['Month']+'月'+dict_old_cases[caseNum]['Date']['Day'] + '日'
                    this_date = f'{med_Date_year}年{med_Date_month}月{med_Date_day}日'

                    workbook = openpyxl.load_workbook(os.path.join(dict_old_cases[caseNum]['Path'],'02.調解事件處理單.xlsx'))
                    workbook.save(os.path.join(record_dir,path_dispal))

                    for read_file_name in readfile_list:
                        doc = Document(os.path.join(dict_old_cases[caseNum]['Path'],read_file_name))
                        for paragraph in doc.paragraphs:
                            if old_case_date in paragraph.text:
                                paragraph.text = paragraph.text.replace(old_case_date,this_date)
                                setKaiTi(paragraph.runs[0].font)
                        # Save the modified document to the same file name
                        doc.save(os.path.join(record_dir,read_file_name))
                    isNewCase = False
                except Exception as e:
                    print(f'發生錯誤:\n{e}\n視為新案處理')
                    

            if isNewCase:    
                manuscript_path = os.path.join(app_dir, path_dispal)
                workbook = openpyxl.load_workbook(manuscript_path)
                sheet = workbook[workbook.sheetnames[0]]
                sheet['D3'].value = f'{receiveNum:03d}'
                sheet['B4'].value = '中華民國' + attendanceTime
                sheet['B6'].value = complete_case_num
                sheet['B7'].value = reason
                sheet['B9'].value = applicant
                sheet['B10'].value = '\n'.join(opponents.values())
                sheet['D9'].value = '身\n行' if '車禍' in reason else '身\n'
                sheet['D10'].value = '\n'.join(['身'] * len(opponents)) if '車禍' in reason else '\n'.join('身' * len(opponents))
                workbook.save(os.path.join(record_dir,path_dispal))
                             
                for read_file_name in readfile_list:
                    manuscript_path = os.path.join(app_dir, read_file_name)
                    doc = Document(manuscript_path)
                    # 開啟 Word 檔案
                    if 'CustomNumberStyle' in doc.styles:
                        doc.styles['CustomNumberStyle'].delete()
                    custom_style = create_custom_style(doc, 'CustomNumberStyle')
                    startFillRow = False
                    # 遍歷每個表格
                    for table in doc.tables:
                        # 遍歷每個行
                        for i, row in enumerate(table.rows):
                            row_texts = []
                            # 針對單一儲存格修改
                            for cell in row.cells:
                                row_texts.append(cell.text)
                                if '收件日期：' in cell.text:
                                    # Replace '收件日期:' with the desired text
                                    cell.text = f'收件日期：{receiveDate}'
                                    setKaiTi(cell.paragraphs[0].runs[0].font, 12, True)
                                    
                                if '收件編號：' in cell.text:
                                    # Replace '收件日期:' with the desired text
                                    if read_file_name == path_request:
                                        cell.text = f'收件編號：{receiveNum:03d}     '
                                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                    else:
                                        cell.text = f'收件編號：{receiveNum:03d}     全1頁'
                                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                                    setKaiTi(cell.paragraphs[0].runs[0].font, 12, True)
                                    
                                    
                                if '案號：'in cell.text and read_file_name == path_request:
                                    cell.text = f'案號{complete_case_num}'
                                    setKaiTi(cell.paragraphs[0].runs[0].font, 12, True)
                                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                                elif '案號'in cell.text and read_file_name != path_request:
                                    cell.text = f'案號{complete_case_num}'
                                    setKaiTi(cell.paragraphs[0].runs[0].font, 12, True)
                                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                                if '上開當事人間因' in cell.text:
                                    cell.text = {f'上開當事人間因「{reason}」聲請調解，於民國{attendanceTime}在本會調解室調解成立，內容如下：'}
                                    setKaiTi(cell.paragraphs[0].runs[0].font, 12, False)
                                #描述事件詳細內容
                                if '聲請人' in cell.text and '自稱' in cell.text and '對造人' in cell.text:
                                    cell.text = '    ' + Description_event
                                    setKaiTi(cell.paragraphs[0].runs[0].font, 16, False)
                                    setFixedRowHeight(cell, 24)
                                    setFixedLineSpacing(cell.paragraphs[0], Pt(24))
                                if '此致' in cell.paragraphs[0].text and '中華民國' in cell.paragraphs[1].text:
                                    cell.paragraphs[1].text = '中華民國' + attendanceTime.split('日')[0] + '日'
                                    run_set_spacing(cell.paragraphs[1].runs[0], 50)
                                    setKaiTi(cell.paragraphs[1].runs[0].font, 14, False)
                                    

                            # 填入兩造資料
                            if row_texts[0] == '稱謂' or row_texts[0] == '稱  謂':
                                startFillRow = True
                            elif startFillRow and row_texts[0] == '聲請人':
                                sexual = main_dict[caseNum+'/'+applicant][1]['性別']
                                birth = str(main_dict[caseNum+'/'+applicant][1]['生日'])
                                id = main_dict[caseNum+'/'+applicant][1]['身份證字號']
                                address = main_dict[caseNum+'/'+applicant][0]['地址']
                                phoneNum = '0'+str(main_dict[caseNum+'/'+applicant][1]['電話號碼'])
                                if read_file_name == path_request:
                                    new_row_texts = ['聲請人', applicant, applicant, sexual, birth, birth, id, '', address, address, phoneNum, phoneNum]
                                elif read_file_name == path_record or read_file_name == path_mediation:
                                    new_row_texts = ['聲請人', applicant, sexual, birth, id, '', address]
                                # 填入新的行資料
                                for new_text, cell in zip(new_row_texts, row.cells):
                                    if new_text:
                                        cell.text = new_text
                                        setKaiTi(cell.paragraphs[0].runs[0].font, 12, False)
                                        if cell.text != address:
                                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                #開始填入對造人
                                itemCount = 1
                                queue_opponents = opponents.copy()
                                while queue_opponents:
                                    opponent = queue_opponents.pop(f'對造人{itemCount}')
                                    sexual = main_dict[caseNum+'/'+opponent][1]['性別']
                                    birth = str(main_dict[caseNum+'/'+opponent][1]['生日'])
                                    id = main_dict[caseNum+'/'+opponent][1]['身份證字號']
                                    address = main_dict[caseNum+'/'+opponent][0]['地址']
                                    phoneNum = '0'+str(main_dict[caseNum+'/'+opponent][1]['電話號碼'])

                                    if read_file_name == path_request: #針對調解聲請書
                                        new_row_texts = ['對造人', opponent , opponent, sexual, birth, birth, id, '', address, address, phoneNum, phoneNum]     
                                    elif read_file_name == path_record or read_file_name == path_mediation: #針對調解筆錄與調解書
                                        new_row_texts = ['對造人', opponent, sexual, birth, id, '', address]

                                    new_row_index = i + itemCount  # 根據itemCount插入位置在当前行的下方
                                    row = addNewRow(table,new_row_index + 2)
                                    # 填入新的行資料
                                    for new_text, cell in zip(new_row_texts, row.cells):
                                        cell.text = new_text
                                        
                                    #合併儲存格並添加框線、文字置中    
                                    for cell in merge_cells_in_row(table, new_row_index).cells:
                                        set_cell_border(cell)
                                        cell.text = cell.text.strip()
                                        setKaiTi(cell.paragraphs[0].runs[0].font, 12, False)
                                        set_cell_border(cell, start={'sz': 2, 'val': 'single'}, end={'sz': 2, 'val': 'single'}, insideV={'sz': 2, 'val': 'single'}, top={'sz': 2, 'val': 'single'})
                                        if cell.text != address:
                                           cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                           
                                    itemCount += 1

                    #遍歷每個段落
                    for i, paragraph in enumerate(doc.paragraphs) :
                        time_att = attendanceTime.split('(')[0] + attendanceTime.split(')')[1] #去除(星期)
                        if '上開當事人間因' in paragraph.text:
                            paragraph.text = f'    上開當事人間因「{reason}」，於民國{time_att}在本會調解室調解成立，內容如下：'
                            setKaiTi(paragraph.runs[0].font,12,False)
                            
                        if '雙方同意如下' in paragraph.text:
                            paragraph.text = f'    {Description_event}，雙方同意如下：'
                            setKaiTi(paragraph.runs[0].font,12,False)
                            if Description_agreement :
                                for j, agreement in enumerate(Description_agreement.split('\n')):
                                    next_paragraph = doc.paragraphs[i+j+1]
                                    bullet_num = agreement.split('、')[0] + '、'
                                    new_paragraph = next_paragraph.insert_paragraph_before(f'{agreement[0]}、', style=custom_style) #往上插入一行
                                    new_paragraph.text += agreement.replace(bullet_num,'')
                                    setKaiTi(new_paragraph.runs[0].font,12,False)
                        # if '主  席：' in  paragraph.text and '紀  錄：' in  paragraph.text:
                        #     paragraph.text = f'主  席：{chairman}         紀  錄：{organiser}'
                        #     setKaiTi(paragraph.runs[0].font,12,True)  

                    #遍歷每個段落
                    for i, paragraph in enumerate(doc.paragraphs) :
                        if '上開調解成立內容：經向當場兩造當事人朗讀或交付閱讀，並無異議。' in paragraph.text:
                            paragraph_to_insert = []
                            this_line = i #從第i行開始
                            this_line += 1 #移動至下一行
                            this_paragraph = doc.paragraphs[this_line] 
                            this_paragraph.text = f"聲請人: {applicant}                                      "  # 聲請人：
                            queue_opponents = opponents.copy()
                            for itemCount in range(1,len(queue_opponents)+1):
                                opponent = queue_opponents.pop(f'對造人{itemCount}')
                                this_paragraph.text += f'對造人: {opponent}'
                                setKaiTi(this_paragraph.runs[0].font,12,False)
                                this_line += 1 #移動至下一行
                                this_paragraph = doc.paragraphs[this_line]
                                if itemCount == 1 :
                                    this_paragraph = this_paragraph.insert_paragraph_before('                〈簽名或蓋章〉                                                                                         〈簽名或蓋章〉')
                                else:
                                    this_paragraph = this_paragraph.insert_paragraph_before('                                                                                                                       〈簽名或蓋章〉') 
                                setKaiTi(this_paragraph.runs[0].font,6,False)
                                setFixedLineSpacing(this_paragraph, Pt(7))  
                                this_line += 1 #移動至下一行
                                this_paragraph = doc.paragraphs[this_line]
                                this_paragraph = this_paragraph.insert_paragraph_before('                                                    ')
                                setKaiTi(this_paragraph.runs[0].font,12,False)

                            this_paragraph.text = '中華民國'+attendanceTime.split('日')[0] + '日'
                            setKaiTi(this_paragraph.runs[0].font,12,False)  
                            if len(opponents) < 2:
                               setFixedLineSpacing(this_paragraph, Pt(50))
                            this_line += 1 #移動至下一行
                            this_paragraph = doc.paragraphs[this_line]
                            write_line = f'                                    主  席：{chairman}            紀  錄：{organiser}'
                            this_paragraph = this_paragraph.insert_paragraph_before(write_line)
                            setKaiTi(this_paragraph.runs[0].font,12,True)   
                            this_line += 1 #移動至下一行
                            this_paragraph = doc.paragraphs[this_line]   
                            this_paragraph = this_paragraph.insert_paragraph_before('                                                                                        〈簽名或蓋章〉                                     〈簽名或蓋章〉')
                            setFixedLineSpacing(this_paragraph, Pt(7))  
                            setKaiTi(this_paragraph.runs[0].font,6,False) 

                    # Save the modified document to the same file name
                    save_path = os.path.join(record_dir,read_file_name)
                    doc.save(save_path)
                    case = Case( receiveNum, caseNum, applicant, opponents, [med_Date_year,med_Date_month,med_Date_day], record_dir).to_dict()
                    dict_old_cases[caseNum] = case
        opponents = {}

    # 将字典保存为 JSON 文件
    with open(os.path.join(dir_path,'oldCaseDict.json'), "w", encoding="utf-8") as file:
        json.dump( dict_old_cases,file,ensure_ascii=False)
except Exception as e:
   print('發生錯誤:')
   traceback.print_exc()
   input(f"請按任意鍵離開")
   sys.exit()
